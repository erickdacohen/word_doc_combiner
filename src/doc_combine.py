import os
import docx
import sys

argv = sys.argv
if argv[1]:
    PATH_DATA_RAW = sys.argv[1]

if argv[2]:
    PATH_DATA_OUTPUT = sys.argv[2]
else:
    PATH_DATA_OUTPUT = "./"


def combine_word_doc(data_raw_path: str = None, data_output_path: str = None) -> int:
    """
    Combines all docx files in a directory into one docx file
    """
    if not data_raw_path:
        print("Missing filepath to raw files")
        return 1

    if not data_output_path:
        print("Warning: No output path specified. Will output to current directory")

    files = os.listdir(data_raw_path)
    output = docx.Document()

    for file in files:
        with open(file) as f:
            text = f.read()
            output.add_paragraph(text)
            output.add_paragraph(" ")

    output.save(str(data_output_path) + "combined_file.docx")

    print("Files combined. Please find result in 'data-output' folder")
    return 0


if __name__ == "__main__":
    combine_word_doc()

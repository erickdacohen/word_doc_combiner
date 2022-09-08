import os
import docx

PATH_DATA_RAW = "./data-raw/"
PATH_DATA_OUTPUT = "./data-output/"

def combine_word_doc(
    data_raw_path = PATH_DATA_RAW, 
    data_output_path = PATH_DATA_OUTPUT
    ):
    
    files = os.listdir(data_raw_path)
    output = docx.Document()

    for file in files:
        with open(file) as f:
            text = f.read()
            output.add_paragraph(text)
            output.add_paragraph(" ")


    output.save(str(data_output_path) + "combined_file.docx")

    print("Files combined. Please find result in 'data-output' folder")

if __name__ == "__main__":
    combine_word_doc()